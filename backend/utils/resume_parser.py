<<<<<<< HEAD
# utils/resume_parser.py
import docx
import PyPDF2

def parse_resume(file):
    text = ""
    try:
        if file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
        elif file.name.endswith(".docx"):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        file.seek(0)

    except Exception as e:
        text = f"⚠️ Error parsing resume: {e}"

=======
# utils/resume_parser.py
import docx
import PyPDF2

def parse_resume(file):
    text = ""
    try:
        if file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
        elif file.name.endswith(".docx"):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        file.seek(0)

    except Exception as e:
        text = f"⚠️ Error parsing resume: {e}"

>>>>>>> 0d47c57f7f5a6c16473e48e4b0355153d2628fc7
    return text.strip()